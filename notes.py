# app.py
# =============================================================================
# CODE ORGANIZER + FILE CONVERTER (STREAMLIT)
# -----------------------------------------------------------------------------
# What this app does:
# 1) Organize (Python-aware): Safely rearranges TOP-LEVEL blocks only
#    (docstring/imports/classes/functions/main-guard/other). Never rewrites inside blocks.
# 2) Convert: Converts common text/data/code files to other formats (safe, offline):
#    - txt/md/py/js/ts -> txt, md, html (basic), pdf, docx
#    - json -> json, yaml (if PyYAML installed), txt
#    - yaml -> yaml, json (if PyYAML installed), txt
#    - csv <-> xlsx (via openpyxl)
#    - xlsx -> csv (first sheet)
# 3) Works with Paste Text or Upload Files + ZIP download output
#
# Mobile-friendly:
# - "Mobile view" toggle forces single-column layout and larger controls.
# =============================================================================

from __future__ import annotations

import io
import os
import re
import csv
import json
import zipfile
import difflib
from dataclasses import dataclass
from typing import List, Tuple, Optional, Dict, Any

import streamlit as st

# Optional libs (graceful degradation)
try:
    import yaml  # type: ignore
    HAS_YAML = True
except Exception:
    HAS_YAML = False

try:
    from docx import Document  # type: ignore
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
    HAS_PDF = True
except Exception:
    HAS_PDF = False

try:
    from openpyxl import Workbook, load_workbook  # type: ignore
    HAS_XLSX = True
except Exception:
    HAS_XLSX = False


# =========================
# Page config + small CSS
# =========================
st.set_page_config(page_title="🧹 Organizer + Converter", page_icon="🧹", layout="wide")

st.markdown(
    """
<style>
/* Slightly bigger tap targets on mobile */
button[kind="primary"], button[kind="secondary"] { min-height: 44px; }
textarea, input { font-size: 14px; }
.small-note { opacity: 0.8; font-size: 12px; }
code { white-space: pre-wrap !important; }
</style>
""",
    unsafe_allow_html=True,
)


# =========================
# Data structures (Organizer)
# =========================
@dataclass
class Block:
    kind: str
    start: int
    end: int
    text: str
    name: str = ""


PY_EXTS = {".py"}
TEXT_EXTS = {".txt", ".md", ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".js", ".ts", ".tsx", ".jsx", ".css", ".html", ".csv", ".xlsx"}
TRIPLE_QUOTE_RE = re.compile(r"^\s*(\"\"\"|''')")


# =========================
# Common helpers
# =========================
def normalize_newlines(s: str) -> str:
    return s.replace("\r\n", "\n").replace("\r", "\n")


def split_lines(content: str) -> List[str]:
    return normalize_newlines(content).split("\n")


def join_lines(lines: List[str]) -> str:
    return "\n".join(lines)


def safe_decode(b: bytes) -> str:
    return b.decode("utf-8", errors="replace")


def safe_encode(s: str) -> bytes:
    return s.encode("utf-8")


def ext_of(name: str) -> str:
    return os.path.splitext(name.lower())[1]


def is_probably_python(filename: str, content: str) -> bool:
    ext = ext_of(filename)
    if ext in PY_EXTS:
        return True
    # heuristic for pasted content / unknown
    c = content
    return ("def " in c) or ("class " in c) or ("import " in c) or ("from " in c)


def count_indent(line: str) -> int:
    s = line.replace("\t", " " * 4)
    return len(s) - len(s.lstrip(" "))


def make_zip(files: Dict[str, bytes]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    return buf.getvalue()


def unified_diff(a: str, b: str, fromfile: str = "original", tofile: str = "organized") -> str:
    a_lines = a.splitlines(keepends=True)
    b_lines = b.splitlines(keepends=True)
    diff = difflib.unified_diff(a_lines, b_lines, fromfile=fromfile, tofile=tofile)
    return "".join(diff)


# =========================
# Organizer: Python block parsing
# =========================
def detect_module_docstring(lines: List[str]) -> Optional[Tuple[int, int]]:
    i = 0
    while i < len(lines) and (lines[i].strip() == "" or lines[i].startswith("#!") or "coding" in lines[i]):
        i += 1
        if i > 15:
            break

    if i >= len(lines):
        return None

    if TRIPLE_QUOTE_RE.match(lines[i]):
        quote = TRIPLE_QUOTE_RE.match(lines[i]).group(1)
        if lines[i].count(quote) >= 2 and lines[i].strip() != quote:
            return (i, i + 1)
        j = i + 1
        while j < len(lines):
            if quote in lines[j]:
                return (i, j + 1)
            j += 1
    return None


def detect_import_block(lines: List[str], start: int) -> Optional[Tuple[int, int]]:
    i = start
    if i >= len(lines):
        return None

    def is_import_line(ln: str) -> bool:
        s = ln.strip()
        return s.startswith("import ") or s.startswith("from ")

    if not is_import_line(lines[i]):
        return None

    j = i
    while j < len(lines):
        s = lines[j].strip()
        if s == "" or s.startswith("#") or is_import_line(lines[j]):
            j += 1
            continue
        break
    return (i, j)


def detect_main_guard(lines: List[str], start: int) -> Optional[Tuple[int, int]]:
    if start >= len(lines):
        return None
    if re.match(r"^\s*if\s+__name__\s*==\s*['\"]__main__['\"]\s*:\s*$", lines[start]):
        i = start + 1
        while i < len(lines) and lines[i].strip() == "":
            i += 1
        base_indent = None
        if i < len(lines):
            base_indent = count_indent(lines[i])
        j = i
        while j < len(lines):
            if lines[j].strip() == "":
                j += 1
                continue
            if base_indent is None:
                break
            if count_indent(lines[j]) < base_indent:
                break
            j += 1
        return (start, j)
    return None


def detect_def_or_class(lines: List[str], start: int) -> Optional[Tuple[int, int, str, str]]:
    if start >= len(lines):
        return None

    line = lines[start]
    if count_indent(line) != 0:
        return None

    m_def = re.match(r"^\s*def\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(.*\)\s*:\s*$", line)
    m_cls = re.match(r"^\s*class\s+([A-Za-z_][A-Za-z0-9_]*)\s*(\(.*\))?\s*:\s*$", line)
    if not (m_def or m_cls):
        return None

    kind = "function" if m_def else "class"
    name = (m_def.group(1) if m_def else m_cls.group(1)) or ""

    # decorators directly above
    i = start
    k = start - 1
    while k >= 0 and lines[k].strip().startswith("@") and count_indent(lines[k]) == 0:
        i = k
        k -= 1

    j = start + 1
    while j < len(lines) and lines[j].strip() == "":
        j += 1
    base_indent = None
    if j < len(lines):
        base_indent = count_indent(lines[j])

    if base_indent is None or base_indent == 0:
        return (i, start + 1, kind, name)

    end = j
    while end < len(lines):
        if lines[end].strip() == "":
            end += 1
            continue
        if count_indent(lines[end]) < base_indent:
            break
        end += 1

    return (i, end, kind, name)


def parse_python_blocks(content: str) -> List[Block]:
    lines = split_lines(content)
    n = len(lines)
    blocks: List[Block] = []
    used = [False] * n

    def mark(s: int, e: int):
        for idx in range(s, e):
            used[idx] = True

    md = detect_module_docstring(lines)
    if md:
        s, e = md
        blocks.append(Block("module_docstring", s, e, join_lines(lines[s:e]).rstrip("\n")))
        mark(s, e)

    i = 0
    while i < n:
        if used[i]:
            i += 1
            continue

        mg = detect_main_guard(lines, i)
        if mg:
            s, e = mg
            blocks.append(Block("main_guard", s, e, join_lines(lines[s:e]).rstrip("\n"), name="__main__"))
            mark(s, e)
            i = e
            continue

        if count_indent(lines[i]) == 0:
            imp = detect_import_block(lines, i)
            if imp:
                s, e = imp
                blocks.append(Block("import", s, e, join_lines(lines[s:e]).rstrip("\n")))
                mark(s, e)
                i = e
                continue

        dc = detect_def_or_class(lines, i)
        if dc:
            s, e, kind, name = dc
            blocks.append(Block(kind, s, e, join_lines(lines[s:e]).rstrip("\n"), name=name))
            mark(s, e)
            i = e
            continue

        # other contiguous region until next known block
        s = i
        j = i + 1
        while j < n and not used[j]:
            if detect_main_guard(lines, j):
                break
            if count_indent(lines[j]) == 0 and detect_import_block(lines, j):
                break
            if detect_def_or_class(lines, j):
                break
            j += 1

        blocks.append(Block("other", s, j, join_lines(lines[s:j]).rstrip("\n")))
        mark(s, j)
        i = j

    blocks.sort(key=lambda b: b.start)
    return blocks


def reorder_python_blocks(blocks: List[Block]) -> Tuple[str, List[str]]:
    report: List[str] = []

    doc = [b for b in blocks if b.kind == "module_docstring"]
    imports = [b for b in blocks if b.kind == "import"]
    main = [b for b in blocks if b.kind == "main_guard"]
    classes = [b for b in blocks if b.kind == "class"]
    funcs = [b for b in blocks if b.kind == "function"]
    others = [b for b in blocks if b.kind == "other"]

    def other_score(b: Block) -> float:
        lines = b.text.split("\n")
        nonempty = [ln for ln in lines if ln.strip() != ""]
        if not nonempty:
            return 0.0
        commentish = sum(1 for ln in nonempty if ln.strip().startswith("#"))
        assignish = sum(1 for ln in nonempty if re.match(r"^\s*[A-Z_][A-Z0-9_]*\s*=", ln))
        return (commentish + assignish) / max(1, len(nonempty))

    early_others, late_others = [], []
    for b in others:
        if b.start < 80 and other_score(b) >= 0.5:
            early_others.append(b)
        else:
            late_others.append(b)

    new_order = doc + imports + early_others + classes + funcs + late_others + main

    original = sorted(blocks, key=lambda x: x.start)
    if [b.start for b in original] != [b.start for b in new_order]:
        report.append("Reordered top-level blocks into: docstring → imports → header/constants → classes → functions → other → main guard.")
        for idx, b in enumerate(new_order):
            label = b.kind + (f" {b.name}" if b.kind in {"class", "function"} and b.name else "")
            report.append(f"{idx+1:02d}. {label} (lines {b.start+1}-{b.end})")
    else:
        report.append("No reordering was necessary (file already matches the target layout).")

    out_parts = []
    for b in new_order:
        txt = b.text.rstrip()
        if txt != "":
            out_parts.append(txt)
    new_text = "\n\n".join(out_parts).rstrip() + "\n"
    return new_text, report


def section_text_generic(content: str) -> Tuple[str, List[str]]:
    lines = split_lines(content)
    report = ["Non-Python mode: no structural reordering performed (safe fallback)."]
    return join_lines(lines).rstrip() + "\n", report


def organize_one(filename: str, content: str, enable_python_reorder: bool) -> Tuple[str, List[str], str]:
    content = normalize_newlines(content).rstrip() + "\n"
    if enable_python_reorder and is_probably_python(filename, content):
        blocks = parse_python_blocks(content)
        organized, report = reorder_python_blocks(blocks)
        diff_text = unified_diff(content, organized, fromfile=f"{filename} (original)", tofile=f"{filename} (organized)")
        return organized, report, diff_text
    organized, report = section_text_generic(content)
    diff_text = unified_diff(content, organized, fromfile=f"{filename} (original)", tofile=f"{filename} (organized)")
    return organized, report, diff_text


# =========================
# Converter
# =========================
def simple_md_to_html(md: str) -> str:
    """
    Basic, dependency-free markdown-ish conversion.
    If you install 'markdown' package, you can upgrade this to real markdown rendering.
    """
    # minimal escaping for safety
    esc = (
        md.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    # very light headings
    lines = esc.split("\n")
    out: List[str] = []
    for ln in lines:
        if ln.startswith("###### "):
            out.append(f"<h6>{ln[7:]}</h6>")
        elif ln.startswith("##### "):
            out.append(f"<h5>{ln[6:]}</h5>")
        elif ln.startswith("#### "):
            out.append(f"<h4>{ln[5:]}</h4>")
        elif ln.startswith("### "):
            out.append(f"<h3>{ln[4:]}</h3>")
        elif ln.startswith("## "):
            out.append(f"<h2>{ln[3:]}</h2>")
        elif ln.startswith("# "):
            out.append(f"<h1>{ln[2:]}</h1>")
        else:
            # preserve line breaks
            out.append(ln)
    body = "<br/>\n".join(out)
    return f"<!doctype html><html><head><meta charset='utf-8'/><meta name='viewport' content='width=device-width,initial-scale=1'/></head><body style='font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial; line-height:1.45; padding:16px;'>{body}</body></html>\n"


def text_to_pdf_bytes(title: str, text: str) -> bytes:
    if not HAS_PDF:
        raise RuntimeError("PDF conversion requires reportlab, but it is not available.")
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    # Basic page layout
    x = 40
    y = height - 50
    c.setTitle(title)

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 22

    c.setFont("Helvetica", 10)
    for ln in normalize_newlines(text).split("\n"):
        # wrap long lines crudely
        while len(ln) > 110:
            chunk, ln = ln[:110], ln[110:]
            c.drawString(x, y, chunk)
            y -= 14
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
        c.drawString(x, y, ln)
        y -= 14
        if y < 50:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 10)

    c.save()
    return buf.getvalue()


def text_to_docx_bytes(title: str, text: str) -> bytes:
    if not HAS_DOCX:
        raise RuntimeError("DOCX conversion requires python-docx, but it is not available.")
    doc = Document()
    doc.add_heading(title, level=1)
    for ln in normalize_newlines(text).split("\n"):
        doc.add_paragraph(ln)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def csv_to_xlsx_bytes(csv_text: str) -> bytes:
    if not HAS_XLSX:
        raise RuntimeError("XLSX conversion requires openpyxl, but it is not available.")
    wb = Workbook()
    ws = wb.active
    reader = csv.reader(io.StringIO(normalize_newlines(csv_text)))
    for r_idx, row in enumerate(reader, start=1):
        for c_idx, val in enumerate(row, start=1):
            ws.cell(row=r_idx, column=c_idx, value=val)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def xlsx_to_csv_bytes(xlsx_bytes: bytes) -> bytes:
    if not HAS_XLSX:
        raise RuntimeError("XLSX conversion requires openpyxl, but it is not available.")
    buf = io.BytesIO(xlsx_bytes)
    wb = load_workbook(buf, data_only=True)
    ws = wb.worksheets[0]  # first sheet
    out = io.StringIO()
    writer = csv.writer(out, lineterminator="\n")
    for row in ws.iter_rows(values_only=True):
        writer.writerow(["" if v is None else v for v in row])
    return out.getvalue().encode("utf-8")


def json_to_yaml_text(obj: Any) -> str:
    if not HAS_YAML:
        raise RuntimeError("YAML conversion requires PyYAML (yaml), but it is not available.")
    return yaml.safe_dump(obj, sort_keys=False, allow_unicode=True)


def yaml_to_obj(text: str) -> Any:
    if not HAS_YAML:
        raise RuntimeError("YAML conversion requires PyYAML (yaml), but it is not available.")
    return yaml.safe_load(text)


def available_targets_for_ext(src_ext: str) -> List[str]:
    src_ext = src_ext.lower()
    targets: List[str] = []
    if src_ext in {".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".css", ".html"}:
        targets += [".txt", ".md", ".html"]
        if HAS_PDF:
            targets += [".pdf"]
        if HAS_DOCX:
            targets += [".docx"]
    elif src_ext == ".json":
        targets += [".json", ".txt"]
        if HAS_YAML:
            targets += [".yaml"]
    elif src_ext in {".yaml", ".yml"}:
        targets += [".yaml", ".txt"]
        if HAS_YAML:
            targets += [".json"]
    elif src_ext == ".csv":
        targets += [".csv", ".txt"]
        if HAS_XLSX:
            targets += [".xlsx"]
    elif src_ext == ".xlsx":
        targets += [".xlsx"]
        targets += [".csv"]  # first sheet
    else:
        # fallback: treat as text
        targets += [".txt"]
        if HAS_PDF:
            targets += [".pdf"]
        if HAS_DOCX:
            targets += [".docx"]
    # unique keep order
    seen = set()
    out = []
    for t in targets:
        if t not in seen:
            out.append(t)
            seen.add(t)
    return out


def convert_file_bytes(filename: str, data: bytes, target_ext: str) -> Tuple[str, bytes, List[str]]:
    """
    Returns (new_filename, new_bytes, report_lines).
    """
    src_ext = ext_of(filename)
    base = os.path.splitext(filename)[0]
    report: List[str] = []

    target_ext = target_ext.lower()
    if not target_ext.startswith("."):
        target_ext = "." + target_ext

    # TEXTY SOURCES
    texty_src = src_ext in {".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".css", ".html"} or src_ext == ""
    if texty_src and target_ext in {".txt", ".md", ".html", ".pdf", ".docx"}:
        text = safe_decode(data)
        if target_ext == ".txt":
            out = normalize_newlines(text).rstrip() + "\n"
            report.append(f"Converted {src_ext or '(unknown)'} → .txt (UTF-8).")
            return f"{base}.txt", safe_encode(out), report

        if target_ext == ".md":
            out = normalize_newlines(text).rstrip() + "\n"
            report.append(f"Converted {src_ext or '(unknown)'} → .md (content preserved as-is).")
            return f"{base}.md", safe_encode(out), report

        if target_ext == ".html":
            if src_ext == ".html":
                out = normalize_newlines(text).rstrip() + "\n"
                report.append("HTML → HTML (normalized newlines).")
                return f"{base}.html", safe_encode(out), report
            if src_ext == ".md":
                html = simple_md_to_html(text)
                report.append("MD → HTML (basic, dependency-free renderer).")
                return f"{base}.html", safe_encode(html), report
            # other text/code -> preformatted HTML
            esc = (
                text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            html = f"<!doctype html><html><head><meta charset='utf-8'/><meta name='viewport' content='width=device-width,initial-scale=1'/></head><body style='font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; padding:16px;'><pre>{esc}</pre></body></html>\n"
            report.append(f"{src_ext or 'text'} → HTML (wrapped in <pre>).")
            return f"{base}.html", safe_encode(html), report

        if target_ext == ".pdf":
            pdf = text_to_pdf_bytes(title=filename, text=text)
            report.append(f"{src_ext or 'text'} → PDF.")
            return f"{base}.pdf", pdf, report

        if target_ext == ".docx":
            docx = text_to_docx_bytes(title=filename, text=text)
            report.append(f"{src_ext or 'text'} → DOCX.")
            return f"{base}.docx", docx, report

    # JSON
    if src_ext == ".json":
        text = safe_decode(data)
        obj = json.loads(text)
        if target_ext == ".json":
            out = json.dumps(obj, indent=2, ensure_ascii=False) + "\n"
            report.append("JSON → JSON (pretty-printed).")
            return f"{base}.json", safe_encode(out), report
        if target_ext in {".yaml", ".yml"}:
            out = json_to_yaml_text(obj)
            report.append("JSON → YAML.")
            return f"{base}.yaml", safe_encode(out), report
        if target_ext == ".txt":
            out = json.dumps(obj, indent=2, ensure_ascii=False) + "\n"
            report.append("JSON → TXT (pretty-printed JSON).")
            return f"{base}.txt", safe_encode(out), report

    # YAML
    if src_ext in {".yaml", ".yml"}:
        text = safe_decode(data)
        obj = yaml_to_obj(text)
        if target_ext in {".yaml", ".yml"}:
            out = json_to_yaml_text(obj)
            report.append("YAML → YAML (normalized).")
            return f"{base}.yaml", safe_encode(out), report
        if target_ext == ".json":
            out = json.dumps(obj, indent=2, ensure_ascii=False) + "\n"
            report.append("YAML → JSON.")
            return f"{base}.json", safe_encode(out), report
        if target_ext == ".txt":
            out = normalize_newlines(text).rstrip() + "\n"
            report.append("YAML → TXT (raw YAML as text).")
            return f"{base}.txt", safe_encode(out), report

    # CSV / XLSX
    if src_ext == ".csv":
        text = safe_decode(data)
        if target_ext == ".csv":
            out = normalize_newlines(text).rstrip() + "\n"
            report.append("CSV → CSV (normalized newlines).")
            return f"{base}.csv", safe_encode(out), report
        if target_ext == ".xlsx":
            xlsx = csv_to_xlsx_bytes(text)
            report.append("CSV → XLSX.")
            return f"{base}.xlsx", xlsx, report
        if target_ext == ".txt":
            out = normalize_newlines(text).rstrip() + "\n"
            report.append("CSV → TXT (raw CSV).")
            return f"{base}.txt", safe_encode(out), report

    if src_ext == ".xlsx":
        if target_ext == ".xlsx":
            report.append("XLSX → XLSX (no-op).")
            return f"{base}.xlsx", data, report
        if target_ext == ".csv":
            out = xlsx_to_csv_bytes(data)
            report.append("XLSX → CSV (first sheet only).")
            return f"{base}.csv", out, report

    # fallback: just emit as txt
    text = safe_decode(data)
    out = normalize_newlines(text).rstrip() + "\n"
    report.append(f"Fallback convert → .txt (source {src_ext or 'unknown'}).")
    return f"{base}.txt", safe_encode(out), report


# =========================
# Sidebar Controls
# =========================
st.title("🧹 Code Organizer + File Converter")

with st.sidebar:
    st.header("Mode")
    app_mode = st.radio("Choose tool", ["Organize", "Convert"], index=0)

    st.divider()
    st.header("Layout")
    mobile_view = st.toggle("Mobile view (single column)", value=False)

    st.divider()
    if app_mode == "Organize":
        st.header("Organizer")
        enable_python_reorder = st.checkbox("Enable Python structural reordering", value=True)
        show_move_report = st.checkbox("Show move report", value=True)
        show_diff = st.checkbox("Show unified diff", value=True)
        st.caption("Organizer is conservative: moves only whole top-level blocks.")
    else:
        st.header("Converter")
        st.caption("Converter is offline and conservative (no risky parsing).")
        if not HAS_YAML:
            st.warning("YAML support: install PyYAML to enable JSON↔YAML conversions.")
        if not HAS_XLSX:
            st.warning("XLSX support: openpyxl missing, CSV↔XLSX disabled.")
        if not HAS_PDF:
            st.warning("PDF support: reportlab missing, TXT→PDF disabled.")
        if not HAS_DOCX:
            st.warning("DOCX support: python-docx missing, TXT→DOCX disabled.")


# =========================
# ORGANIZE UI
# =========================
if app_mode == "Organize":
    st.caption("Paste text or upload files, then organize safely. Download single file or ZIP.")

    input_mode = st.segmented_control("Input", options=["Paste text", "Upload files"], default="Paste text")

    if input_mode == "Paste text":
        if mobile_view:
            st.subheader("Original")
            filename = st.text_input("Filename (helps detection)", value="pasted.py")
            raw = st.text_area("Paste your text/code here", height=320)

            run = st.button("Organize", type="primary", use_container_width=True, disabled=(raw.strip() == ""))

            if "org_paste_out" not in st.session_state:
                st.session_state.org_paste_out = ""
                st.session_state.org_paste_report = []
                st.session_state.org_paste_diff = ""

            if run:
                organized, report, diff_text = organize_one(filename, raw, enable_python_reorder)
                st.session_state.org_paste_out = organized
                st.session_state.org_paste_report = report
                st.session_state.org_paste_diff = diff_text

            st.subheader("Organized")
            if st.session_state.org_paste_out:
                st.code(st.session_state.org_paste_out, language="python" if filename.endswith(".py") else "text")
                st.download_button(
                    "Download organized file",
                    data=st.session_state.org_paste_out.encode("utf-8"),
                    file_name=filename,
                    mime="text/plain",
                    use_container_width=True,
                )
                if show_move_report:
                    with st.expander("Move report", expanded=False):
                        st.write("\n".join([f"- {x}" for x in st.session_state.org_paste_report]))
                if show_diff:
                    with st.expander("Unified diff", expanded=False):
                        st.code(st.session_state.org_paste_diff, language="diff")
            else:
                st.info("Paste content and press **Organize**.")
        else:
            colA, colB = st.columns([1, 1], gap="large")
            with colA:
                st.subheader("Original")
                filename = st.text_input("Filename (helps detection)", value="pasted.py")
                raw = st.text_area("Paste your text/code here", height=420)
                run = st.button("Organize", type="primary", use_container_width=True, disabled=(raw.strip() == ""))

            if "org_paste_out" not in st.session_state:
                st.session_state.org_paste_out = ""
                st.session_state.org_paste_report = []
                st.session_state.org_paste_diff = ""

            if run:
                organized, report, diff_text = organize_one(filename, raw, enable_python_reorder)
                st.session_state.org_paste_out = organized
                st.session_state.org_paste_report = report
                st.session_state.org_paste_diff = diff_text

            with colB:
                st.subheader("Organized")
                if st.session_state.org_paste_out:
                    st.code(st.session_state.org_paste_out, language="python" if filename.endswith(".py") else "text")
                    st.download_button(
                        "Download organized file",
                        data=st.session_state.org_paste_out.encode("utf-8"),
                        file_name=filename,
                        mime="text/plain",
                        use_container_width=True,
                    )
                else:
                    st.info("Paste content on the left and click **Organize**.")

            if show_move_report and st.session_state.org_paste_report:
                with st.expander("Move report", expanded=False):
                    st.write("\n".join([f"- {x}" for x in st.session_state.org_paste_report]))

            if show_diff and st.session_state.org_paste_diff:
                with st.expander("Unified diff", expanded=False):
                    st.code(st.session_state.org_paste_diff, language="diff")

    else:
        uploaded = st.file_uploader(
            "Upload one or more files",
            type=sorted({e.lstrip(".") for e in (PY_EXTS | TEXT_EXTS)}),
            accept_multiple_files=True,
        )

        if not uploaded:
            st.info("Upload files to begin.")
        else:
            st.write(f"Files loaded: **{len(uploaded)}**")
            process = st.button("Organize all files", type="primary", use_container_width=True)

            if process:
                results: Dict[str, bytes] = {}
                per_file_reports: Dict[str, List[str]] = {}
                per_file_diffs: Dict[str, str] = {}

                for f in uploaded:
                    name = f.name
                    raw = safe_decode(f.getvalue())
                    organized, report, diff_text = organize_one(name, raw, enable_python_reorder)
                    results[name] = organized.encode("utf-8")
                    per_file_reports[name] = report
                    per_file_diffs[name] = diff_text

                st.success("Done.")

                st.subheader("Preview")
                for name in results:
                    with st.expander(name, expanded=False):
                        is_py = name.lower().endswith(".py")
                        st.code(results[name].decode("utf-8", errors="replace"), language="python" if is_py else "text")
                        st.download_button(
                            f"Download {name}",
                            data=results[name],
                            file_name=name,
                            mime="text/plain",
                            key=f"org_dl_{name}",
                        )
                        if show_move_report:
                            st.markdown("**Move report:**")
                            st.write("\n".join([f"- {x}" for x in per_file_reports[name]]))
                        if show_diff:
                            st.markdown("**Diff:**")
                            st.code(per_file_diffs[name], language="diff")

                zip_bytes = make_zip(results)
                st.download_button(
                    "Download ALL organized files (ZIP)",
                    data=zip_bytes,
                    file_name="organized_files.zip",
                    mime="application/zip",
                    use_container_width=True,
                )

# =========================
# CONVERT UI
# =========================
else:
    st.caption("Upload files and convert them to another format. Download single file(s) or a ZIP bundle.")

    uploaded = st.file_uploader(
        "Upload one or more files",
        type=sorted({e.lstrip(".") for e in (TEXT_EXTS | PY_EXTS | {".zip"})}),
        accept_multiple_files=True,
    )

    if not uploaded:
        st.info("Upload files to convert.")
    else:
        st.write(f"Files loaded: **{len(uploaded)}**")

        # If user uploads a zip, offer to extract and convert contents (simple)
        contains_zip = any(ext_of(f.name) == ".zip" for f in uploaded)
        if contains_zip:
            st.warning("ZIP uploaded: this app can extract and convert contained files (non-binary) in a simple way.")

        # Build per-file target selectors
        st.subheader("Conversion settings")
        file_targets: Dict[str, str] = {}

        for f in uploaded:
            name = f.name
            e = ext_of(name)
            # ZIP: treat specially
            if e == ".zip":
                continue
            targets = available_targets_for_ext(e)
            default = targets[0] if targets else ".txt"
            file_targets[name] = st.selectbox(
                f"Target for {name}",
                options=targets,
                index=targets.index(default) if default in targets else 0,
                key=f"target_{name}",
            )

        st.divider()

        # Optional: organize python before conversion (useful for .py -> .pdf/.docx/.html etc.)
        pre_organize = st.checkbox("Organize Python files before converting (safe reorder)", value=False)
        enable_python_reorder_for_convert = True  # always safe; gated by checkbox

        convert_btn = st.button("Convert", type="primary", use_container_width=True)

        if convert_btn:
            out_files: Dict[str, bytes] = {}
            out_reports: Dict[str, List[str]] = {}

            def convert_one(name: str, data: bytes, tgt: str):
                # optionally organize python first if it's text-based conversion
                if pre_organize:
                    try:
                        txt = safe_decode(data)
                        if is_probably_python(name, txt):
                            organized, _, _ = organize_one(name, txt, enable_python_reorder_for_convert)
                            data = organized.encode("utf-8")
                    except Exception:
                        # if decode fails, skip organize
                        pass

                new_name, new_bytes, report = convert_file_bytes(name, data, tgt)
                out_files[new_name] = new_bytes
                out_reports[new_name] = report

            for f in uploaded:
                name = f.name
                e = ext_of(name)
                raw = f.getvalue()

                if e == ".zip":
                    # Extract zip, convert each file if it looks text-like, skip large binaries
                    try:
                        z = zipfile.ZipFile(io.BytesIO(raw))
                        for zi in z.infolist():
                            if zi.is_dir():
                                continue
                            inner_name = zi.filename
                            inner_ext = ext_of(inner_name)
                            if inner_ext not in (TEXT_EXTS | PY_EXTS):
                                continue
                            data = z.read(zi)
                            # pick a target based on extension
                            targets = available_targets_for_ext(inner_ext)
                            tgt = targets[0] if targets else ".txt"
                            convert_one(inner_name, data, tgt)
                    except Exception as ex:
                        st.error(f"Could not read zip {name}: {ex}")
                    continue

                tgt = file_targets.get(name, ".txt")
                try:
                    convert_one(name, raw, tgt)
                except Exception as ex:
                    st.error(f"Failed to convert {name} → {tgt}: {ex}")

            if out_files:
                st.success("Conversion complete.")

                st.subheader("Outputs")
                for name in out_files:
                    with st.expander(name, expanded=False):
                        ext = ext_of(name)
                        if ext in {".txt", ".md", ".json", ".yaml", ".yml", ".csv", ".html", ".py", ".js", ".ts", ".tsx", ".jsx"}:
                            st.code(safe_decode(out_files[name]), language="text")
                        else:
                            st.caption("Binary output (preview not shown).")

                        st.download_button(
                            f"Download {name}",
                            data=out_files[name],
                            file_name=name,
                            mime="application/octet-stream",
                            key=f"conv_dl_{name}",
                        )
                        st.markdown("**Report:**")
                        st.write("\n".join([f"- {x}" for x in out_reports.get(name, [])]))

                zip_bytes = make_zip(out_files)
                st.download_button(
                    "Download ALL converted files (ZIP)",
                    data=zip_bytes,
                    file_name="converted_files.zip",
                    mime="application/zip",
                    use_container_width=True,
                )
            else:
                st.warning("No output files were produced (nothing convertible or conversion failed).")

st.markdown("<div class='small-note'>Tip: Want more conversions (e.g., DOCX↔PDF, HTML↔MD, or full Markdown rendering)? Tell me which formats you care about and I’ll extend the converter safely.</div>", unsafe_allow_html=True)
